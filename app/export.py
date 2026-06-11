"""Export the draft report markdown to .docx."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def md_to_docx(md: str, out_path: str | Path) -> Path:
    doc = Document()
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").strip()) <= {"-", " "}:
            # markdown table
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                if not set("".join(cells)) <= {"-", " "}:
                    rows.append(cells)
                i += 1
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        table.cell(r, c).text = re.sub(r"\*\*", "", cell)
                        for para in table.cell(r, c).paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(8)
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            doc.add_heading(re.sub(r"\*\*", "", m.group(2)), level=len(m.group(1)))
        elif line.startswith(">"):
            p = doc.add_paragraph(re.sub(r"\*\*|^>\s*", "", line))
            p.italic = True
        elif line.strip():
            doc.add_paragraph(re.sub(r"\*\*", "", line))
        i += 1
    out = Path(out_path)
    doc.save(out)
    return out


def main(argv: list[str]) -> int:
    src = Path(argv[0] if argv else "data/seeds/mcintyre/draft_report.md")
    out = src.with_suffix(".docx")
    md_to_docx(src.read_text(), out)
    print(f"written {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
